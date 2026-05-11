from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import argparse
from datetime import datetime
from typing import List, Tuple, Optional
import sys
from pathlib import Path


class TxtToDocxConverter:
    def __init__(self, 
                 source_folder: str = None,
                 output_file: str = None,
                 include_subfolders: bool = True,
                 add_toc: bool = False,
                 add_header_footer: bool = False,
                 font_name: str = "宋体",
                 font_size: int = 12,
                 page_size: str = "A4",
                 line_spacing: float = 1.5):
        """
        初始化转换器
        
        Args:
            source_folder: 源文件夹路径
            output_file: 输出文件路径
            include_subfolders: 是否包含子文件夹
            add_toc: 是否添加目录
            add_header_footer: 是否添加页眉页脚
            font_name: 字体名称
            font_size: 字体大小
            page_size: 页面大小 (A4, Letter, etc.)
            line_spacing: 行间距
        """
        self.source_folder = Path(source_folder) if source_folder else None
        self.output_file = Path(output_file) if output_file else None
        self.include_subfolders = include_subfolders
        self.add_toc = add_toc
        self.add_header_footer = add_header_footer
        self.font_name = font_name
        self.font_size = font_size
        self.page_size = page_size
        self.line_spacing = line_spacing
        
        # 支持的页面大小
        self.page_sizes = {
            'A4': (Inches(8.27), Inches(11.69)),
            'Letter': (Inches(8.5), Inches(11)),
            'A3': (Inches(11.69), Inches(16.54)),
            'B5': (Inches(6.93), Inches(9.84)),
        }
        
        # 统计信息
        self.stats = {
            'total_files': 0,
            'processed_files': 0,
            'total_pages': 0,
            'start_time': datetime.now(),
            'file_list': []
        }
    
    def validate_paths(self) -> Tuple[bool, str]:
        """验证路径"""
        if not self.source_folder or not self.source_folder.exists():
            return False, f"源文件夹不存在: {self.source_folder}"
        
        if not self.source_folder.is_dir():
            return False, f"源路径不是文件夹: {self.source_folder}"
        
        # 设置默认输出文件名
        if not self.output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            self.output_file = self.source_folder.parent / f"合并文档_{timestamp}.docx"
        
        # 确保输出目录存在
        self.output_file.parent.mkdir(parents=True, exist_ok=True)
        
        return True, "验证通过"
    
    def find_txt_files(self) -> List[Path]:
        """查找所有txt文件"""
        txt_files = []
        
        if self.include_subfolders:
            pattern = "**/*.tsx"
        else:
            pattern = "*.tsx"
        
        for txt_file in self.source_folder.glob(pattern):
            if txt_file.is_file() and txt_file.suffix.lower() == '.tsx':
                txt_files.append(txt_file)
        
        # 按文件名排序
        txt_files.sort(key=lambda x: x.name.lower())
        
        self.stats['total_files'] = len(txt_files)
        self.stats['file_list'] = [str(f.relative_to(self.source_folder)) for f in txt_files]
        
        return txt_files
    
    def read_file_content(self, file_path: Path) -> Tuple[bool, str, str]:
        """读取文件内容，自动检测编码"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                return True, content, encoding
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return False, "", str(e)
        
        return False, "", "无法解码文件"
    
    def create_document_styles(self, doc: Document):
        """创建文档样式"""
        # 设置页面大小
        if self.page_size in self.page_sizes:
            width, height = self.page_sizes[self.page_size]
            section = doc.sections[0]
            section.page_width = width
            section.page_height = height
        
        # 创建标题样式
        styles = doc.styles
        
        # 标题1样式
        title1 = styles.add_style('CustomTitle1', WD_STYLE_TYPE.PARAGRAPH)
        title1.font.name = '黑体'
        title1.font.size = Pt(22)
        title1.font.bold = True
        title1.font.color.rgb = RGBColor(0, 0, 0)
        title1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title1.paragraph_format.space_after = Pt(24)
        
        # 标题2样式（用于文件名）
        title2 = styles.add_style('CustomTitle2', WD_STYLE_TYPE.PARAGRAPH)
        title2.font.name = self.font_name
        title2.font.size = Pt(16)
        title2.font.bold = True
        title2.font.color.rgb = RGBColor(0, 0, 139)  # 深蓝色
        title2.paragraph_format.space_before = Pt(18)
        title2.paragraph_format.space_after = Pt(12)
        
        # 正文样式
        normal = styles['Normal']
        normal.font.name = self.font_name
        normal.font.size = Pt(self.font_size)
        normal.paragraph_format.line_spacing = self.line_spacing
        normal.paragraph_format.space_after = Pt(6)
        
        # 代码样式（如果检测到可能是代码）
        code_style = styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(10)
        code_style.font.color.rgb = RGBColor(0, 100, 0)  # 深绿色
        code_style.paragraph_format.left_indent = Inches(0.5)
    
    def add_header_footer_content(self, doc: Document):
        """添加页眉页脚"""
        if not self.add_header_footer:
            return
        
        for section in doc.sections:
            # 页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"合并文档 - {self.source_folder.name}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 页脚
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"第 \\p 页 / 共 \\n 页"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_table_of_contents(self, doc: Document, file_list: List[str]):
        """添加目录"""
        if not self.add_toc or not file_list:
            return
        
        # 添加目录标题
        toc_title = doc.add_paragraph("目 录")
        toc_title.style = 'CustomTitle1'
        
        # 添加目录内容
        for i, filename in enumerate(file_list, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {filename}")
            
            # 添加制表符和页码占位符
            run.add_tab()
            run.add_text("......")
            run.add_tab()
            run.add_text("1")  # Word会自动更新页码
        
        # 添加分页符
        doc.add_page_break()
    
    def process_file_content(self, content: str, file_path: Path) -> List[str]:
        """处理文件内容，智能分段"""
        lines = content.split('\n')
        paragraphs = []
        current_para = []
        
        for line in lines:
            line = line.rstrip('\r')
            
            # 空行处理：如果遇到空行，结束当前段落
            if line.strip() == '':
                if current_para:
                    paragraphs.append('\n'.join(current_para))
                    current_para = []
            else:
                current_para.append(line)
        
        # 添加最后一个段落
        if current_para:
            paragraphs.append('\n'.join(current_para))
        
        # 如果段落太少，考虑更细的分段
        if len(paragraphs) < 3 and len(lines) > 10:
            # 可能是代码或特殊格式，按原样保留
            return [content]
        
        return paragraphs
    
    def is_likely_code(self, content: str) -> bool:
        """判断内容是否可能是代码"""
        code_indicators = ['import ', 'def ', 'class ', 'function ', 
                          '{', '}', ';', '//', '/*', '*/', '#include']
        
        lines = content.split('\n')[:20]  # 检查前20行
        code_line_count = 0
        
        for line in lines:
            if any(indicator in line for indicator in code_indicators):
                code_line_count += 1
        
        return code_line_count >= 3
    
    def convert(self) -> Tuple[bool, str]:
        """执行转换"""
        # 验证路径
        is_valid, message = self.validate_paths()
        if not is_valid:
            return False, message
        
        # 查找文件
        txt_files = self.find_txt_files()
        if not txt_files:
            return False, f"在 {self.source_folder} 中未找到txt文件"
        
        print(f"找到 {len(txt_files)} 个txt文件")
        
        # 创建文档
        doc = Document()
        
        # 创建样式
        self.create_document_styles(doc)
        
        # 添加页眉页脚
        self.add_header_footer_content(doc)
        
        # 添加目录
        if self.add_toc:
            self.add_table_of_contents(doc, self.stats['file_list'])
        
        # 添加文档标题
        title = doc.add_paragraph(f"文档合并报告")
        title.style = 'CustomTitle1'
        
        # 添加文档信息
        info = doc.add_paragraph()
        info.add_run(f"源文件夹: {self.source_folder}\n")
        info.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info.add_run(f"文件数量: {len(txt_files)}")
        
        doc.add_page_break()
        
        # 处理每个文件
        for idx, txt_file in enumerate(txt_files, 1):
            print(f"处理文件 {idx}/{len(txt_files)}: {txt_file.name}")
            
            try:
                # 读取文件
                success, content, encoding = self.read_file_content(txt_file)
                if not success:
                    print(f"  警告: 无法读取文件 {txt_file.name}, 编码: {encoding}")
                    continue
                
                # 添加文件名作为标题
                file_title = doc.add_paragraph(f"{idx}. {txt_file.name}")
                file_title.style = 'CustomTitle2'
                
                # 添加文件信息
                file_info = doc.add_paragraph()
                file_info.add_run(f"路径: {txt_file.relative_to(self.source_folder)}\n")
                file_info.add_run(f"编码: {encoding}\n")
                file_info.add_run(f"大小: {os.path.getsize(txt_file)} 字节")
                
                # 处理并添加内容
                paragraphs = self.process_file_content(content, txt_file)
                
                # 判断内容类型并应用相应样式
                is_code = self.is_likely_code(content)
                
                for para_text in paragraphs:
                    if is_code:
                        p = doc.add_paragraph(para_text, style='CodeStyle')
                    else:
                        p = doc.add_paragraph(para_text)
                
                # 添加分页符（最后一个文件除外）
                if idx < len(txt_files):
                    doc.add_page_break()
                
                self.stats['processed_files'] += 1
                
            except Exception as e:
                print(f"  错误: 处理文件 {txt_file.name} 时出错: {str(e)}")
                continue
        
        # 添加统计信息
        doc.add_page_break()
        stats_title = doc.add_paragraph("转换统计")
        stats_title.style = 'CustomTitle2'
        
        end_time = datetime.now()
        duration = (end_time - self.stats['start_time']).total_seconds()
        
        stats_content = doc.add_paragraph()
        stats_content.add_run(f"总文件数: {self.stats['total_files']}\n")
        stats_content.add_run(f"成功处理: {self.stats['processed_files']}\n")
        stats_content.add_run(f"失败文件: {self.stats['total_files'] - self.stats['processed_files']}\n")
        stats_content.add_run(f"耗时: {duration:.2f} 秒\n")
        stats_content.add_run(f"输出文件: {self.output_file}")
        
        # 保存文档
        try:
            doc.save(str(self.output_file))
            print(f"文档已保存: {self.output_file}")
            
            # 尝试打开文档（可选）
            if sys.platform == 'win32':
                os.startfile(self.output_file)
            elif sys.platform == 'darwin':
                os.system(f'open "{self.output_file}"')
            
            return True, f"转换完成，共处理 {self.stats['processed_files']} 个文件"
            
        except Exception as e:
            return False, f"保存文档时出错: {str(e)}"


def parse_arguments():
    """解析命令行参数"""
    parser = argparse.ArgumentParser(description='将多个TXT文件合并为单个DOCX文档')
    
    parser.add_argument('source', nargs='?', help='源文件夹路径')
    parser.add_argument('-o', '--output', help='输出文件路径')
    parser.add_argument('--no-subfolders', action='store_true', 
                       help='不包含子文件夹')
    parser.add_argument('--toc', action='store_true', 
                       help='添加目录')
    parser.add_argument('--header-footer', action='store_true',
                       help='添加页眉页脚')
    parser.add_argument('--font', default='宋体',
                       help='字体名称 (默认: 宋体)')
    parser.add_argument('--size', type=int, default=12,
                       help='字体大小 (默认: 12)')
    parser.add_argument('--page', default='A4',
                       choices=['A4', 'Letter', 'A3', 'B5'],
                       help='页面大小 (默认: A4)')
    parser.add_argument('--spacing', type=float, default=1.5,
                       help='行间距 (默认: 1.5)')
    parser.add_argument('--list', action='store_true',
                       help='仅列出文件，不转换')
    
    return parser.parse_args()


def main():
    """主函数"""
    args = parse_arguments()
    
    # 如果没有提供源文件夹，交互式询问
    if not args.source:
        args.source = input("请输入源文件夹路径: ").strip()
    
    if not args.source:
        print("错误: 必须提供源文件夹路径")
        return
    
    # 创建转换器
    converter = TxtToDocxConverter(
        source_folder=args.source,
        output_file=args.output,
        include_subfolders=not args.no_subfolders,
        add_toc=args.toc,
        add_header_footer=args.header_footer,
        font_name=args.font,
        font_size=args.size,
        page_size=args.page,
        line_spacing=args.spacing
    )
    
    # 如果仅列出文件
    if args.list:
        txt_files = converter.find_txt_files()
        print(f"找到 {len(txt_files)} 个txt文件:")
        for i, file in enumerate(txt_files, 1):
            print(f"  {i:3d}. {file.relative_to(converter.source_folder)}")
        return
    
    # 执行转换
    success, message = converter.convert()
    
    if success:
        print(f"✓ {message}")
    else:
        print(f"✗ {message}")
        sys.exit(1)


if __name__ == '__main__':
    main()
